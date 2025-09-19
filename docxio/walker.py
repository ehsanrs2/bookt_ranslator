"""
Walk the DOCX document tree and yield translatable text units.
"""

from __future__ import annotations

from dataclasses import dataclass
from typing import Iterable, Iterator, List, Optional, Sequence, Tuple

from docx.document import Document as _DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from utils.filters import is_field_code_paragraph, is_field_code_oxml


@dataclass
class TextUnit:
    """A single translatable unit: a paragraph with runs.

    Depending on origin, either `paragraph` is set (python-docx object), or
    `p_oxml` is set (lxml CT_P element for shapes/textboxes).
    """

    location: str
    text: str
    runs: Sequence[Run] | Sequence[object]
    paragraph: Optional[Paragraph] = None
    p_oxml: Optional[object] = None
    context: str = "body"


class DocxWalker:
    def __init__(
        self,
        document: _DocxDocument,
        *,
        include_headers: bool = True,
        include_footers: bool = True,
        include_shapes: bool = True,
        skip_fields: bool = True,
    ) -> None:
        self.document = document
        self.include_headers = include_headers
        self.include_footers = include_footers
        self.include_shapes = include_shapes
        self.skip_fields = skip_fields

    def iter_units(self) -> Iterator[TextUnit]:
        # Body paragraphs
        for idx, p in enumerate(self.document.paragraphs, start=1):
            if not p.text:
                continue
            if self.skip_fields and is_field_code_paragraph(p):
                continue
            yield TextUnit(
                location=f"body:p[{idx}]",
                text=p.text,
                runs=list(p.runs),
                paragraph=p,
                context="body",
            )

        # Tables in body
        for t_index, table in enumerate(self.document.tables, start=1):
            yield from self._iter_table(table, prefix=f"body:table[{t_index}]", context="body")

        # Headers/Footers
        for s_idx, section in enumerate(self.document.sections, start=1):
            if self.include_headers:
                try:
                    header = section.header
                    for p_idx, p in enumerate(header.paragraphs, start=1):
                        if not p.text:
                            continue
                        if self.skip_fields and is_field_code_paragraph(p):
                            continue
                        yield TextUnit(
                            location=f"header[{s_idx}]:p[{p_idx}]",
                            text=p.text,
                            runs=list(p.runs),
                            paragraph=p,
                            context="header",
                        )
                    for t_idx, table in enumerate(header.tables, start=1):
                        yield from self._iter_table(
                            table, prefix=f"header[{s_idx}]:table[{t_idx}]", context="header"
                        )
                except RecursionError:
                    # Some documents cause recursion with linked headers; skip gracefully
                    pass
                except Exception:
                    pass
            if self.include_footers:
                try:
                    footer = section.footer
                    for p_idx, p in enumerate(footer.paragraphs, start=1):
                        if not p.text:
                            continue
                        if self.skip_fields and is_field_code_paragraph(p):
                            continue
                        yield TextUnit(
                            location=f"footer[{s_idx}]:p[{p_idx}]",
                            text=p.text,
                            runs=list(p.runs),
                            paragraph=p,
                            context="footer",
                        )
                    for t_idx, table in enumerate(footer.tables, start=1):
                        yield from self._iter_table(
                            table, prefix=f"footer[{s_idx}]:table[{t_idx}]", context="footer"
                        )
                except RecursionError:
                    pass
                except Exception:
                    pass

        # Text boxes / shapes
        if self.include_shapes:
            yield from self._iter_shapes()

    def _iter_table(self, table: Table, *, prefix: str, context: str) -> Iterator[TextUnit]:
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                loc = f"{prefix}/cell[{r_idx},{c_idx}]"
                yield from self._iter_cell(cell, loc, context)

    def _iter_cell(self, cell: _Cell, prefix: str, context: str) -> Iterator[TextUnit]:
        for p_idx, p in enumerate(cell.paragraphs, start=1):
            if not p.text:
                continue
            if is_field_code_paragraph(p):
                continue
            yield TextUnit(
                location=f"{prefix}/p[{p_idx}]",
                text=p.text,
                runs=list(p.runs),
                paragraph=p,
                context=context,
            )
        for t_idx, table in enumerate(cell.tables, start=1):
            yield from self._iter_table(
                table, prefix=f"{prefix}/table[{t_idx}]", context=context
            )

    def _iter_shapes(self) -> Iterator[TextUnit]:
        # Find shapes' text box paragraphs via lxml on the document element
        root = self.document.element
        p_nodes = root.xpath('.//w:txbxContent//w:p')
        for s_idx, p in enumerate(p_nodes, start=1):
            # skip field code paragraphs
            if self.skip_fields and is_field_code_oxml(p):
                continue
            # gather runs and text
            r_nodes = p.xpath('.//w:r')
            t_nodes = p.xpath('.//w:r//w:t')
            text = "".join((t.text or "") for t in t_nodes)
            if not text.strip():
                continue
            yield TextUnit(
                location=f"shape:p[{s_idx}]",
                text=text,
                runs=list(r_nodes),
                p_oxml=p,
                context="shape",
            )
