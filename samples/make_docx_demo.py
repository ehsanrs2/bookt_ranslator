#!/usr/bin/env python3
"""
Generate a tiny DOCX sample with paragraphs, bold/italic, list, table, header/footer, and a text box.
"""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


def add_text_box(paragraph, text: str):
    # Creating proper text boxes via raw XML is verbose; keep sample simple.
    # For compatibility, just append a labeled paragraph acting as a "shape" placeholder.
    run = paragraph.add_run(f"[Textbox: {text}]")


def main():
    doc = Document()
    # Header & footer
    hdr = doc.sections[0].header
    hdr_p = hdr.add_paragraph("Demo header with number ")
    hdr_p.add_run("123").bold = True
    ftr = doc.sections[0].footer
    ftr.add_paragraph("Page ")

    p1 = doc.add_paragraph("Hello world from a sample DOCX.")
    p1.runs[0].bold = True
    p2 = doc.add_paragraph()
    p2.add_run("This is ")
    r = p2.add_run("italic")
    r.italic = True
    p2.add_run(" and ")
    r2 = p2.add_run("bold")
    r2.bold = True
    p2.add_run(" text.")

    # Hyperlink-like text (plain URL)
    doc.add_paragraph("Visit https://example.com for more info.")

    # List
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")

    # Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Language"
    table.cell(1, 1).text = "English/French"

    # Pseudo text box marker
    tp = doc.add_paragraph("Text box:")
    add_text_box(tp, "Text box content")

    out = "samples/demo.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
